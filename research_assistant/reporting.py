"""Render a finished research task to Markdown (default), DOCX, or PDF.

Shared by every frontend: the CLI writes the bytes to `exports/<slug>.<ext>`,
the Telegram bot sends them as an attachment — one renderer, many sinks.

The report body from the pipeline is Markdown, so DOCX/PDF are rendered from a
deliberately small Markdown subset — headings (`#`/`##`/`###`), paragraphs and
bullet lists — which is all the agents actually emit. That keeps us off a heavy
document toolchain (no pandoc/LaTeX/GTK): `python-docx` and `fpdf2` are pure-pip
optional extras (the `export` extra), and a missing one becomes a one-line
install hint, not a traceback.

`task` is a plain dict with keys: status, id, query, final_report, sources,
total_tokens — so this module stays free of any storage/ORM dependency.
"""

from __future__ import annotations

import os
import re
from datetime import datetime
from io import BytesIO
from pathlib import Path

FORMATS = ("md", "docx", "pdf", "tex", "paper")

# file extension per format (default: the format name itself). "paper" is a
# tectonic-compiled LaTeX document — the artifact is a PDF.
_EXT = {"paper": "pdf"}


def ext_for(fmt: str) -> str:
    """Filename extension for a format ('paper' renders to a .pdf)."""
    return _EXT.get(fmt, fmt)


def slugify(text: str, maxlen: int = 60) -> str:
    """Filesystem-safe stem derived from the query: 'Who is Ronaldo?' -> 'who-is-ronaldo'."""
    s = re.sub(r"[^\w\s-]", "", text.lower())
    s = re.sub(r"[\s_-]+", "-", s).strip("-")
    return s[:maxlen].strip("-") or "report"


def render(task: dict, fmt: str = "md") -> bytes:
    """The report in `fmt` as bytes (md is UTF-8 text; docx/pdf are binary).

    Raises ValueError for an unknown format, and ModuleNotFoundError/RuntimeError
    (with a friendly message) when a format's optional dependency or a required
    Unicode font is missing. Does not check task status — callers gate on that."""
    if fmt == "md":
        return _markdown(task).encode("utf-8")
    if fmt == "docx":
        return _docx_bytes(task)
    if fmt == "pdf":
        return _pdf_bytes(task)
    if fmt in ("tex", "paper"):
        # lazy: LaTeX/APA export lives in its own module (latex.py)
        from research_assistant import latex

        return latex.to_tex(task).encode("utf-8") if fmt == "tex" else latex.to_pdf(task)
    raise ValueError(f"unknown format {fmt!r} — choose from {', '.join(FORMATS)}")


def save_report(task: dict, fmt: str = "md") -> Path | None:
    """Write the report in `fmt` to exports/<slug>.<ext>; return the path or None.

    Returns None for an unfinished/empty task. Propagates render()'s errors."""
    if task.get("status") != "done" or not task.get("final_report"):
        return None
    data = render(task, fmt)  # validates fmt / deps before we touch the disk
    exports = Path("exports")
    exports.mkdir(exist_ok=True)
    path = exports / f"{slugify(task.get('query', 'report'))}.{ext_for(fmt)}"
    path.write_bytes(data)
    return path


# --- shared content ----------------------------------------------------------


def _markdown(task: dict) -> str:
    """The full report as a Markdown string (the .md export, verbatim body)."""
    lines = [
        f"# Research report — {task.get('query', '')}",
        "",
        f"*task {str(task.get('id', ''))[:8]} · {datetime.now():%Y-%m-%d %H:%M}*",
        "",
        task.get("final_report", ""),
    ]
    sources = task.get("sources") or []
    if sources:
        lines += ["", "## Sources", ""]
        for i, s in enumerate(sources, 1):
            lines.append(f"{i}. [{s.get('title', '')}]({s.get('url', '')})")
    total = task.get("total_tokens") or 0
    if total:
        lines += ["", f"*tokens: {total:,}*"]
    return "\n".join(lines) + "\n"


def _doc_blocks(task: dict) -> list[tuple[str, str]]:
    """Flatten the report into (kind, text) blocks for the DOCX/PDF emitters.

    kind is one of: title, meta, h1/h2/h3, li, p. The body is parsed from the
    Markdown subset via _blocks; the shell (title, sources, tokens) is added
    around it so both emitters share one structure."""
    blocks: list[tuple[str, str]] = [
        ("title", f"Research report — {task.get('query', '')}"),
        ("meta", f"task {str(task.get('id', ''))[:8]} · {datetime.now():%Y-%m-%d %H:%M}"),
    ]
    blocks += _blocks(task.get("final_report", ""))
    sources = task.get("sources") or []
    if sources:
        blocks.append(("h2", "Sources"))
        for i, s in enumerate(sources, 1):
            blocks.append(("li", f"{i}. {s.get('title', '')} ({s.get('url', '')})"))
    total = task.get("total_tokens") or 0
    if total:
        blocks.append(("meta", f"tokens: {total:,}"))
    return blocks


def _blocks(md: str) -> list[tuple[str, str]]:
    """Parse a Markdown subset into (kind, text) blocks: h1/h2/h3, li, p.

    Consecutive non-blank, non-heading, non-bullet lines fold into one paragraph;
    a blank line ends it. Inline markup is flattened by _inline."""
    out: list[tuple[str, str]] = []
    para: list[str] = []

    def flush() -> None:
        if para:
            out.append(("p", _inline(" ".join(para))))
            para.clear()

    for raw in md.splitlines():
        line = raw.strip()
        if not line:
            flush()
            continue
        h = re.match(r"(#{1,3})\s+(.*)", line)
        if h:
            flush()
            out.append((f"h{len(h.group(1))}", _inline(h.group(2))))
            continue
        b = re.match(r"[-*]\s+(.*)", line)
        if b:
            flush()
            out.append(("li", _inline(b.group(1))))
            continue
        para.append(line)
    flush()
    return out


def _inline(text: str) -> str:
    """Flatten inline Markdown to plain text: links -> 'text (url)', drop emphasis
    and code markers. The document formats carry structure via block kinds, so we
    don't try to preserve bold/italic runs — plain, readable text is enough."""
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)  # [t](u) -> t (u)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)               # **bold**
    text = re.sub(r"(?<!\*)\*([^*]+)\*(?!\*)", r"\1", text)      # *italic*
    text = re.sub(r"`([^`]+)`", r"\1", text)                     # `code`
    return text


def _missing(fmt: str, pkg: str) -> ModuleNotFoundError:
    return ModuleNotFoundError(
        f"{fmt} export needs the '{pkg}' package. Install the export extra: "
        f'pip install "research-assistant[export]"  (or: pip install {pkg})'
    )


# --- DOCX --------------------------------------------------------------------


def _docx_bytes(task: dict) -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt
    except ModuleNotFoundError as e:
        raise _missing("docx", "python-docx") from e

    doc = Document()
    for kind, text in _doc_blocks(task):
        if kind == "title":
            doc.add_heading(text, level=0)
        elif kind in ("h1", "h2", "h3"):
            doc.add_heading(text, level=int(kind[1]))
        elif kind == "li":
            doc.add_paragraph(text, style="List Bullet")
        elif kind == "meta":
            run = doc.add_paragraph().add_run(text)
            run.italic = True
            run.font.size = Pt(9)
        else:  # p
            doc.add_paragraph(text)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# --- PDF ---------------------------------------------------------------------

# Point sizes per block kind (unlisted kinds fall back to body size).
_PDF_SIZES = {"title": 18, "h1": 15, "h2": 13, "h3": 12, "meta": 9}
_PDF_BODY = 11


def _pdf_bytes(task: dict) -> bytes:
    try:
        from fpdf import FPDF
    except ModuleNotFoundError as e:
        raise _missing("pdf", "fpdf2") from e

    blocks = _doc_blocks(task)
    font = _unicode_font()
    # fpdf2's built-in fonts are latin-1 only. Rather than crash mid-render on a
    # Cyrillic (or other non-latin) report, refuse up front with a clear steer.
    if font is None and any(ord(c) > 0xFF for _, t in blocks for c in t):
        raise RuntimeError(
            "PDF export needs a Unicode TTF font for this text and none was found "
            "on this system — use the docx or md format instead."
        )

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    if font is not None:
        pdf.add_font("body", "", font)
        family = "body"
    else:
        family = "helvetica"

    for kind, text in blocks:
        size = _PDF_SIZES.get(kind, _PDF_BODY)
        pdf.set_font(family, size=size)
        pdf.multi_cell(0, size * 0.55, ("•  " + text) if kind == "li" else text)
        pdf.ln(2)
    return bytes(pdf.output())


def _unicode_font() -> str | None:
    """Path to a Unicode-capable TTF on this system, or None. We probe common OS
    fonts rather than bundling one — Arial on Windows, DejaVu/Arial on Linux/mac
    all cover Cyrillic — so the repo stays free of binary assets."""
    candidates = [
        Path(os.environ.get("WINDIR", "C:/Windows")) / "Fonts" / "arial.ttf",
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
        Path("/System/Library/Fonts/Supplemental/Arial.ttf"),
        Path("/Library/Fonts/Arial.ttf"),
    ]
    return next((str(c) for c in candidates if c.is_file()), None)
